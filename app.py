import streamlit as st
import streamlit.components.v1 as components
from docxtpl import DocxTemplate
from io import BytesIO
import csv
import datetime
import os
import requests
import base64
import io
import pandas as pd
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import subprocess

# Беттің баптаулары
st.set_page_config(page_title="Smart Paper Generator", page_icon="📝", layout="wide")

# Сессия күйлерін бастау
if "lang" not in st.session_state:
    st.session_state.lang = "kz"
if "theme" not in st.session_state:
    st.session_state.theme = "light"
if "is_registered" not in st.session_state:
    st.session_state.is_registered = False
if "ui_font" not in st.session_state:
    st.session_state.ui_font = "System Default"

# Динамические счетчики для строк Рисунков и Таблиц
if "fig_count" not in st.session_state:
    st.session_state.fig_count = 1
if "tab_count" not in st.session_state:
    st.session_state.tab_count = 1

# Аудармалар сөздігі
locales = {
    "ru": {
        "title": "📝 Умный генератор научных статей",
        "subtitle": "Вестник ЕНУ им. Л.Н. Гумилева · Химия / География · 2025",
        "btn_theme_dark": "🌙 Тёмная тема",
        "btn_theme_light": "☀️ Светлая тема",
        "nav_gen": "📄 Генератор статей",
        "nav_reg": "👤 Регистрация",
        "sidebar_title": "⚙️ Настройки",
        "lbl_ui_font": "Шрифт интерфейса",
        "lbl_lang": "Основной язык статьи",
        "lbl_sec": "Секция",
        "lbl_type": "Тип статьи",
        "lbl_mrnti": "МРНТИ / IRSTI",
        "sec_meta": "1. Основные метаданные",
        "lbl_title": "Название статьи",
        "lbl_authors": "Авторы",
        "lbl_authors_help": "Например: Имя Фамилия1, Имя Фамилия2",
        "lbl_affil": "Аффилиации (Место работы/учебы)",
        "lbl_affil_help": "1 Университет, Город, Страна; email",
        "lbl_email": "Email для корреспонденции",
        "sec_text": "2. Текст статьи (Загрузка IMRAD)",
        "lbl_abstract": "Аннотация (до 300 слов)",
        "lbl_kw": "Ключевые слова",
        "lbl_kw_help": "Слово 1; слово 2; слово 3 (от 3 до 10 слов)",
        "lbl_intro": "Введение (.txt/.docx)",
        "lbl_methods": "Материалы и методы (.txt/.docx)",
        "lbl_results": "Результаты (.txt/.docx)",
        "lbl_discussion": "Обсуждение (.txt/.docx)",
        "lbl_conclusion": "Заключение (.txt/.docx)",
        "lbl_ref_manager": "📚 Менеджер литературы",
        "lbl_ref_style": "Стиль цитирования",
        "lbl_fig_manager": "📊 Менеджер рисунков",
        "lbl_tab_manager": "📋 Менеджер таблиц",
        "lbl_add_fig": "➕ Добавить рисунок",
        "lbl_add_tab": "➕ Добавить таблицу",
        "btn_upload_short": "📎 Загрузить файл",
        "lbl_fig_hint_title": "💡 Подсказка для сложных графиков",
        "lbl_fig_hint_text": "Если рисунок состоит из нескольких частей (a, b, c), используйте **один тег** `[@fig1]` для всей группы. В подписи подробно опишите каждую часть: *Рисунок 1. Главное название: (a) первый график; (b) второй график.*",
        "lbl_tab_hint_title": "💡 Подсказка для сложных таблиц",
        "lbl_tab_hint_text": "Если ваша таблица очень широкая или имеет объединенные ячейки (merge), мы настоятельно рекомендуем загружать её в формате **.docx**, чтобы сохранить всё форматирование. Обязательно вставьте тег `[@tab1]` в нужный абзац вашего текста.",
        "lbl_samples": "📥 Скачать шаблоны файлов",
        "sec_backmatter": "4. Дополнительная информация (Back Matter)",
        "lbl_supp": "6. Supplementary Materials",
        "lbl_contrib": "7. Author Contributions",
        "lbl_auth_info": "8. Author Information",
        "lbl_funding": "9. Funding",
        "lbl_ack": "10. Acknowledgements",
        "lbl_coi": "11. Conflicts of Interest",
        "sec_trans": "3. Переводы метаданных",
        "trans_info": "По требованиям журнала необходимо предоставить название, авторов, аннотацию и ключевые слова на двух других языках.",
        "gen_btn": "🚀 Сгенерировать статью",
        "err_abs_len": "⚠️ Аннотация слишком длинная: {count} слов. Максимум: 300.",
        "succ_abs_len": "Слов в аннотации: {count}/300",
        "err_fill_req": "Пожалуйста, заполните хотя бы Название и Авторов.",
        "err_gen": "Произошла ошибка при генерации: ",
        "succ_gen": "✅ Документ успешно сгенерирован!",
        "btn_dl_docx": "⬇️ Скачать .docx",
        "btn_dl_pdf": "⬇️ Скачать .pdf",
        "err_pdf": "⚠️ Не удалось сгенерировать PDF (требуется LibreOffice на сервере). Доступен DOCX файл.",
        "reg_header": "📝 Регистрация исследователя",
        "reg_name": "ФИО (Полностью)",
        "reg_email": "Ваш Email",
        "reg_phone": "Номер телефона",
        "reg_org": "Организация / Университет",
        "reg_pos": "Должность / Статус (например: Докторант)",
        "reg_submit": "Зарегистрироваться",
        "reg_success": "✅ Вы успешно зарегистрированы! Теперь вам доступен генератор статей.",
        "reg_info": "Вы можете перейти в раздел «Генератор статей».",
        "reg_req_msg": "🔒 Для создания статьи необходимо заполнить форму регистрации. Перейдите во вкладку «Регистрация» выше.",
        "reg_err_fill": "Пожалуйста, заполните Имя, Email и Телефон.",
        "f_author": "Канат Самарханов / Kanat Samarkhanov",
        "f_license": "Лицензия",
        "f_univ": "ЕНУ им. Л.Н. Гумилева",
        "browse_files": "Выберите файл или перетащите его сюда",
        "drag_drop": "Поддерживаемые форматы: txt, docx, png, jpg, xlsx, csv",
        "limit": "Лимит 200MB",
        "fig_prefix": "Рисунок",
        "tab_prefix": "Таблица"
    },
    "kz": {
        "title": "📝 Ғылыми мақалалардың ақылды генераторы",
        "subtitle": "Л.Н. Гумилев атындағы ЕҰУ Хабаршысы · Химия / География · 2025",
        "btn_theme_dark": "🌙 Түнгі режим",
        "btn_theme_light": "☀️ Күндізгі режим",
        "nav_gen": "📄 Мақала генераторы",
        "nav_reg": "👤 Тіркелу",
        "sidebar_title": "⚙️ Баптаулар",
        "lbl_ui_font": "Интерфейс қаріпі",
        "lbl_lang": "Мақаланың негізгі тілі",
        "lbl_sec": "Секция",
        "lbl_type": "Мақала түрі",
        "lbl_mrnti": "МРНТИ / IRSTI",
        "sec_meta": "1. Негізгі метадеректер",
        "lbl_title": "Мақаланың атауы",
        "lbl_authors": "Авторлар",
        "lbl_authors_help": "Мысалы: Аты Жөні1, Аты Жөні2",
        "lbl_affil": "Аффилиация (Жұмыс/оқу орны)",
        "lbl_affil_help": "1 Университет, Қала, Ел; email",
        "lbl_email": "Корреспонденцияға арналған email",
        "sec_text": "2. Мақала мәтіні (IMRAD Файлдары)",
        "lbl_abstract": "Аңдатпа (300 сөзге дейін)",
        "lbl_kw": "Түйінді сөздер",
        "lbl_kw_help": "Сөз 1; сөз 2; сөз 3 (3-тен 10 сөзге дейін)",
        "lbl_intro": "Кіріспе (.txt/.docx)",
        "lbl_methods": "Материалдар/әдістер (.txt/.docx)",
        "lbl_results": "Нәтижелер (.txt/.docx)",
        "lbl_discussion": "Талқылау (.txt/.docx)",
        "lbl_conclusion": "Қорытынды (.txt/.docx)",
        "lbl_ref_manager": "📚 Әдебиеттер менеджері",
        "lbl_ref_style": "Дәйексөз стилі",
        "lbl_fig_manager": "📊 Суреттер менеджері",
        "lbl_tab_manager": "📋 Кестелер менеджері",
        "lbl_add_fig": "➕ Сурет қосу",
        "lbl_add_tab": "➕ Кесте қосу",
        "btn_upload_short": "📎 Файлды жүктеу",
        "lbl_fig_hint_title": "💡 Күрделі суреттерге арналған нұсқаулық",
        "lbl_fig_hint_text": "Егер сурет бірнеше бөліктен (a, b, c) тұрса, бүкіл топ үшін **бір тегті** `[@fig1]` пайдаланыңыз. Әр бөлікті сипаттаңыз: *Сурет 1. Негізгі атау: (a) бірінші график; (b) екінші график.*",
        "lbl_tab_hint_title": "💡 Күрделі кестелерге арналған нұсқаулық",
        "lbl_tab_hint_text": "Кестеңіз өте кең болса немесе біріктірілген ұяшықтары болса, пішімдеуді сақтау үшін оны **.docx** форматында жүктеуді ұсынамыз. `[@tab1]` тегін мәтініңізге қосуды ұмытпаңыз.",
        "lbl_samples": "📥 Файл үлгілерін жүктеп алу",
        "sec_backmatter": "4. Қосымша ақпарат (Back Matter)",
        "lbl_supp": "6. Supplementary Materials",
        "lbl_contrib": "7. Author Contributions",
        "lbl_auth_info": "8. Author Information",
        "lbl_funding": "9. Funding",
        "lbl_ack": "10. Acknowledgements",
        "lbl_coi": "11. Conflicts of Interest",
        "sec_trans": "3. Метадеректер аудармасы",
        "trans_info": "Журнал талаптарына сәйкес атауын, авторларын, аңдатпасын және түйінді сөздерін басқа екі тілде ұсыну қажет.",
        "gen_btn": "🚀 Мақаланы генерациялау",
        "err_abs_len": "⚠️ Аңдатпа тым ұзын: {count} сөз. Максимум: 300.",
        "succ_abs_len": "Аңдатпадағы сөз саны: {count}/300",
        "err_fill_req": "Кем дегенде Атауын және Авторларын толтырыңыз.",
        "err_gen": "Генерация кезінде қате пайда болды: ",
        "succ_gen": "✅ Құжат сәтті генерацияланды!",
        "btn_dl_docx": "⬇️ .docx жүктеу",
        "btn_dl_pdf": "⬇️ .pdf жүктеу",
        "err_pdf": "⚠️ PDF жасау мүмкін болмады (серверде LibreOffice қажет). Тек DOCX файлы қолжетімді.",
        "reg_header": "📝 Зерттеушіні тіркеу",
        "reg_name": "Аты-жөні (Толық)",
        "reg_email": "Сіздің Email",
        "reg_phone": "Телефон нөмірі",
        "reg_org": "Ұйым / Университет",
        "reg_pos": "Қызметі / Мәртебесі (мысалы: Докторант)",
        "reg_submit": "Тіркелу",
        "reg_success": "✅ Сіз жүйеге сәтті тіркелдіңіз! Енді мақала генераторы қолжетімді.",
        "reg_info": "Сіз «Мақала генераторы» бөліміне өтіп, мақала жасай аласыз.",
        "reg_req_msg": "🔒 Мақала жасау үшін тіркелу формасын толтыру қажет. Жоғарыдағы «Тіркелу» бөліміне өтіңіз.",
        "reg_err_fill": "Аты-жөні, Email және Телефонды толтырыңыз.",
        "f_author": "Канат Самарханов / Kanat Samarkhanov",
        "f_license": "Лицензия",
        "f_univ": "Л.Н. Гумилев атындағы ЕҰУ",
        "browse_files": "Файлды таңдаңыз немесе осында сүйреңіз",
        "drag_drop": "Қолдау көрсетілетін форматтар: txt, docx, png, jpg, xlsx, csv",
        "limit": "Шектеу 200MB",
        "fig_prefix": "Сурет",
        "tab_prefix": "Кесте"
    },
    "en": {
        "title": "📝 Smart Paper Generator",
        "subtitle": "L.N. Gumilyov ENU Bulletin · Chemistry / Geography · 2025",
        "btn_theme_dark": "🌙 Dark mode",
        "btn_theme_light": "☀️ Light mode",
        "nav_gen": "📄 Paper Generator",
        "nav_reg": "👤 Registration",
        "sidebar_title": "⚙️ Paper Settings",
        "lbl_ui_font": "Interface Font",
        "lbl_lang": "Primary Language",
        "lbl_sec": "Section",
        "lbl_type": "Paper Type",
        "lbl_mrnti": "IRSTI / МРНТИ",
        "sec_meta": "1. Basic Metadata",
        "lbl_title": "Article Title",
        "lbl_authors": "Authors",
        "lbl_authors_help": "E.g.: Firstname Lastname1, Firstname Lastname2",
        "lbl_affil": "Affiliations",
        "lbl_affil_help": "1 University, City, Country; email",
        "lbl_email": "Correspondence Email",
        "sec_text": "2. Main Text (IMRAD Uploads)",
        "lbl_abstract": "Abstract (up to 300 words)",
        "lbl_kw": "Keywords",
        "lbl_kw_help": "Keyword 1; keyword 2; keyword 3 (3 to 10 words)",
        "lbl_intro": "Introduction (.txt/.docx)",
        "lbl_methods": "Materials & Methods (.txt/.docx)",
        "lbl_results": "Results (.txt/.docx)",
        "lbl_discussion": "Discussion (.txt/.docx)",
        "lbl_conclusion": "Conclusion (.txt/.docx)",
        "lbl_ref_manager": "📚 Reference Manager",
        "lbl_ref_style": "Citation Style",
        "lbl_fig_manager": "📊 Figure Manager",
        "lbl_tab_manager": "📋 Table Manager",
        "lbl_add_fig": "➕ Add Figure",
        "lbl_add_tab": "➕ Add Table",
        "btn_upload_short": "📎 Upload File",
        "lbl_fig_hint_title": "💡 Hint for Complicated Figures",
        "lbl_fig_hint_text": "If a figure has multiple parts (a, b, c), use a **single tag** `[@fig1]` for the entire group. Describe each part in the caption: *Figure 1. Main title: (a) chart one; (b) chart two.*",
        "lbl_tab_hint_title": "💡 Hint for Complex Tables",
        "lbl_tab_hint_text": "For wide tables or tables with merged cells, please upload a **.docx** file to preserve the formatting perfectly. Make sure to insert the tag `[@tab1]` in your text.",
        "lbl_samples": "📥 Download Sample Files",
        "sec_backmatter": "4. Additional Information (Back Matter)",
        "lbl_supp": "6. Supplementary Materials",
        "lbl_contrib": "7. Author Contributions",
        "lbl_auth_info": "8. Author Information",
        "lbl_funding": "9. Funding",
        "lbl_ack": "10. Acknowledgements",
        "lbl_coi": "11. Conflicts of Interest",
        "sec_trans": "3. Metadata Translations",
        "trans_info": "According to the journal requirements, the title, authors, abstract and keywords must be provided in two other languages.",
        "gen_btn": "🚀 Generate Document",
        "err_abs_len": "⚠️ Abstract is too long: {count} words. Maximum: 300.",
        "succ_abs_len": "Words in abstract: {count}/300",
        "err_fill_req": "Please fill in at least the Title and Authors.",
        "err_gen": "An error occurred during generation: ",
        "succ_gen": "✅ Document successfully generated!",
        "btn_dl_docx": "⬇️ Download .docx",
        "btn_dl_pdf": "⬇️ Download .pdf",
        "err_pdf": "⚠️ Failed to generate PDF (requires LibreOffice on the server). DOCX is available.",
        "reg_header": "📝 Researcher Registration",
        "reg_name": "Full Name",
        "reg_email": "Your Email",
        "reg_phone": "Phone Number",
        "reg_org": "Organization / University",
        "reg_pos": "Position / Status (e.g., PhD Student)",
        "reg_submit": "Register",
        "reg_success": "✅ You have successfully registered! The paper generator is now unlocked.",
        "reg_info": "You can now go to the 'Paper Generator' section.",
        "reg_req_msg": "🔒 To generate an article, you must complete the registration form. Please go to the 'Registration' tab above.",
        "reg_err_fill": "Please fill in your Name, Email, and Phone.",
        "f_author": "Kanat Samarkhanov",
        "f_license": "License",
        "f_univ": "L.N. Gumilyov ENU",
        "browse_files": "Choose a file or drag and drop it here",
        "drag_drop": "Supported formats: txt, docx, png, jpg, xlsx, csv",
        "limit": "Limit 200MB",
        "fig_prefix": "Figure",
        "tab_prefix": "Table"
    }
}

l = locales[st.session_state.lang]

# Dynamic Font Family mapping
font_mapping = {
    "System Default": "sans-serif",
    "Times New Roman": "'Times New Roman', Times, serif",
    "Arial": "Arial, Helvetica, sans-serif",
    "Georgia": "Georgia, serif"
}
selected_css_font = font_mapping.get(st.session_state.ui_font, "sans-serif")

# ------------ CSS Дизайн ------------
# 1. Main File Uploader design (for IMRAD)
# 2. COMPACT File Uploader design (for dynamic rows)
file_uploader_i18n = f"""
<style>
/* Font family and Text Justify distribution */
* {{ font-family: {selected_css_font} !important; }}
.stApp p, .stApp div[data-testid="stMarkdownContainer"] {{ text-align: justify !important; }}

/* Customize MAIN File Uploaders to look like the Screenshot */
[data-testid="stFileUploadDropzone"] {{
    border: 2px dashed #a0aec0 !important; border-radius: 12px !important;
    padding: 24px !important; text-align: center !important;
    display: flex; flex-direction: column; align-items: center; justify-content: center;
}}
[data-testid="stFileUploadDropzone"]::before {{ content: "☁️"; font-size: 40px; display: block; margin-bottom: 8px; }}
[data-testid="stFileUploadDropzone"] svg {{ display: none !important; }}
[data-testid="stFileUploadDropzone"] button {{ 
    color: transparent !important; position: relative; background-color: transparent !important;
    border: none !important; box-shadow: none !important; margin-top: 10px;
}}
[data-testid="stFileUploadDropzone"] button::after {{
    content: "{l['browse_files']}"; color: #3b82f6 !important; position: absolute;
    left: 50%; top: 50%; transform: translate(-50%, -50%); visibility: visible; 
    font-weight: 600; font-size: 16px; white-space: nowrap; text-decoration: underline;
}}
[data-testid="stFileUploadDropzone"] button:hover::after {{ color: #2563eb !important; }}
[data-testid="stFileUploadDropzone"] div[data-testid="stText"] span {{ display: none !important; }}
[data-testid="stFileUploadDropzone"] div[data-testid="stText"]::before {{
    content: "{l['drag_drop']}\\A {l['limit']}"; white-space: pre-wrap; color: #64748b !important; 
    display: block; text-align: center; font-size: 0.85rem; margin-top: 5px;
}}

/* УБИРАЕМ ВЕРТИКАЛЬНЫЕ ЛИНИИ В СЕЛЕКТАХ */
[data-baseweb="select"] input {{ caret-color: transparent !important; }}
div[data-baseweb="select"] > div > div:nth-child(2) {{ width: 0 !important; display: none !important; border: none !important; }}
div[data-baseweb="select"] div[aria-hidden="true"] {{ background-color: transparent !important; width: 0 !important; border: none !important; display: none !important; }}
div[data-baseweb="select"] * {{ border-left: none !important; border-right: none !important; }}

/* --- COMPACT UPLOADER FOR ROWS (Fixing the "Mess") --- */
.element-container:has(.small-uploader-marker) + .element-container [data-testid="stFileUploadDropzone"] {{
    padding: 0 !important; min-height: 40px !important; height: 40px !important;
    border: 1px dashed #94a3b8 !important; border-radius: 6px !important;
    display: flex; align-items: center; justify-content: center; flex-direction: row;
}}
.element-container:has(.small-uploader-marker) + .element-container [data-testid="stFileUploadDropzone"]::before {{ display: none !important; }}
.element-container:has(.small-uploader-marker) + .element-container [data-testid="stFileUploadDropzone"] div[data-testid="stText"] {{ display: none !important; }}
.element-container:has(.small-uploader-marker) + .element-container [data-testid="stFileUploadDropzone"] button {{ margin: 0 !important; width: 100% !important; height: 100% !important; display: block; }}
.element-container:has(.small-uploader-marker) + .element-container [data-testid="stFileUploadDropzone"] button::after {{
    content: "{l['btn_upload_short']}" !important; font-size: 13px !important; color: #64748b !important; text-decoration: none !important;
}}
.element-container:has(.small-uploader-marker) + .element-container [data-testid="stFileUploadDropzone"]:hover {{ border-color: #3b82f6 !important; }}
.element-container:has(.small-uploader-marker) + .element-container [data-testid="stFileUploadDropzone"]:hover button::after {{ color: #3b82f6 !important; }}
</style>
"""

light_css = """
<style>
.stApp { background-color: #ffffff !important; }
[data-testid="stSidebar"] { background-color: #f8f9fa !important; border-right: 1px solid #e9ecef !important; }
[data-testid="stMarkdownContainer"] h1, [data-testid="stMarkdownContainer"] h2, [data-testid="stMarkdownContainer"] h3 { color: #1a3a5c !important; }
hr { border-color: #e9ecef !important; }
input, textarea, [data-baseweb="select"] > div { background-color: #eaf4fc !important; color: #1a3a5c !important; border: 1px solid #bcdcfa !important; border-radius: 6px !important; }
input:focus, textarea:focus, [data-baseweb="select"] > div:focus-within { border-color: #58a6ff !important; box-shadow: 0 0 0 2px rgba(88, 166, 255, 0.2) !important; }
input[disabled], textarea[disabled], [data-baseweb="select"] > div[aria-disabled="true"] { background-color: #e9ecef !important; color: #6c757d !important; -webkit-text-fill-color: #6c757d !important; border: 1px solid #dddddd !important; }
button[kind="primary"] { background-color: #2563eb !important; color: #ffffff !important; border: 1px solid #1d4ed8 !important; border-radius: 6px !important; font-weight: 600 !important; }
button[kind="primary"]:hover { background-color: #1d4ed8 !important; border-color: #1e40af !important; box-shadow: 0 0 8px rgba(37, 99, 235, 0.4) !important; }

/* Segmented Control */
div[data-testid="stRadio"] { display: flex; justify-content: center; margin-bottom: 1rem; }
div[data-testid="stRadio"] div[role="radiogroup"] { background-color: #f1f3f4 !important; border-radius: 20px !important; padding: 4px !important; display: inline-flex !important; gap: 4px !important; border: none !important; }
div[data-testid="stRadio"] div[role="radiogroup"] label { background-color: transparent !important; padding: 8px 24px !important; border-radius: 16px !important; color: #5f6368 !important; font-weight: 500 !important; cursor: pointer !important; border: none !important; transition: all 0.2s; margin:0 !important; }
div[data-testid="stRadio"] div[role="radiogroup"] label:hover { background-color: rgba(0,0,0,0.05) !important; }
div[data-testid="stRadio"] div[role="radio"] { display: none !important; }
div[data-testid="stRadio"] div[role="radiogroup"] label:has(div[aria-checked="true"]) { background-color: #ffffff !important; color: #1a1a1a !important; box-shadow: 0 2px 5px rgba(0,0,0,0.1) !important; font-weight: 600 !important; }

/* Custom light theme uploader dropzone */
[data-testid="stFileUploadDropzone"] { background-color: #f8fafc !important; }
[data-testid="stFileUploadDropzone"]:hover { border-color: #3b82f6 !important; background-color: #eff6ff !important; }
</style>
"""

dark_css = """
<style>
.stApp { background-color: #0b1426 !important; }
[data-testid="stSidebar"] { background-color: #0f1c34 !important; border-right: 1px solid #1d3354 !important; }
[data-testid="stMarkdownContainer"] h1, [data-testid="stMarkdownContainer"] h2, [data-testid="stMarkdownContainer"] h3 { color: #f8fafc !important; }
p, span, label { color: #cbd5e1 !important; }
hr { border-color: #1d3354 !important; }
input, textarea, [data-baseweb="select"] > div { background-color: #132440 !important; color: #f8fafc !important; border: 1px solid #284470 !important; box-shadow: 0 0 4px rgba(46, 92, 184, 0.2) !important; border-radius: 6px !important; }
input:focus, textarea:focus, [data-baseweb="select"] > div:focus-within { border: 1px solid #3b82f6 !important; box-shadow: 0 0 6px rgba(59, 130, 246, 0.6) !important; }
input[disabled], textarea[disabled], [data-baseweb="select"] > div[aria-disabled="true"] { background-color: #0f1c34 !important; color: #64748b !important; -webkit-text-fill-color: #64748b !important; border: 1px solid #1d3354 !important; box-shadow: none !important; }
button[kind="primary"] { background-color: #3b82f6 !important; color: #ffffff !important; border: 1px solid #2563eb !important; border-radius: 6px !important; font-weight: 600 !important; }
button[kind="primary"]:hover { background-color: #60a5fa !important; border-color: #3b82f6 !important; box-shadow: 0 0 8px rgba(96, 165, 250, 0.4) !important; }
button[kind="secondary"] { background-color: #132440 !important; color: #cbd5e1 !important; border: 1px solid #284470 !important; }
button[kind="secondary"]:hover { border-color: #3b82f6 !important; color: #ffffff !important; }

/* Segmented Control */
div[data-testid="stRadio"] { display: flex; justify-content: center; margin-bottom: 1rem; }
div[data-testid="stRadio"] div[role="radiogroup"] { background-color: #0f1c34 !important; border-radius: 20px !important; padding: 4px !important; display: inline-flex !important; gap: 4px !important; border: 1px solid #1d3354 !important; }
div[data-testid="stRadio"] div[role="radiogroup"] label { background-color: transparent !important; padding: 8px 24px !important; border-radius: 16px !important; color: #64748b !important; font-weight: 500 !important; cursor: pointer !important; border: none !important; transition: all 0.2s; margin:0 !important; }
div[data-testid="stRadio"] div[role="radiogroup"] label:hover { background-color: rgba(255,255,255,0.05) !important; color: #cbd5e1 !important;}
div[data-testid="stRadio"] div[role="radio"] { display: none !important; }
div[data-testid="stRadio"] div[role="radiogroup"] label:has(div[aria-checked="true"]) { background-color: #2563eb !important; color: #ffffff !important; box-shadow: 0 2px 5px rgba(0,0,0,0.3) !important; font-weight: 600 !important; }

/* Custom dark theme uploader dropzone */
[data-testid="stFileUploadDropzone"] { background-color: #132440 !important; border-color: #284470 !important; }
[data-testid="stFileUploadDropzone"]:hover { border-color: #3b82f6 !important; background-color: #1d3354 !important; }
</style>
"""

st.markdown(file_uploader_i18n, unsafe_allow_html=True)
st.markdown(dark_css if st.session_state.theme == "dark" else light_css, unsafe_allow_html=True)

# ------------ Helpers ------------
def extract_text(uploaded_file):
    if not uploaded_file: return ""
    try:
        if uploaded_file.name.endswith('.txt'): return uploaded_file.read().decode('utf-8')
        elif uploaded_file.name.endswith('.docx'):
            doc_file = docx.Document(uploaded_file)
            return '\n'.join([p.text for p in doc_file.paragraphs])
    except Exception as e: return f"[Error: {str(e)}]"
    return ""

def create_sample_docx(section_title):
    doc = docx.Document()
    heading = doc.add_heading(section_title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"Here is sample content for {section_title}. Delete this and paste your text. ")
    p.add_run("All paragraphs here are justified. ").bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2 = doc.add_paragraph("Example of tagging: The results shown in [@fig1] are summarized in [@tab1]. Relevant literature supports this [@ref1].")
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def append_to_github_csv(filename, row_data, header_data):
    try:
        github_token = st.secrets["GITHUB_TOKEN"]
        github_repo = st.secrets["GITHUB_REPO"]
    except Exception:
        file_exists = os.path.isfile(filename)
        with open(filename, mode="a", encoding="utf-8-sig", newline="") as f:
            writer = csv.writer(f)
            if not file_exists: writer.writerow(header_data)
            writer.writerow(row_data)
        return
    url = f"https://api.github.com/repos/{github_repo}/contents/{filename}"
    headers = {"Authorization": f"token {github_token}"}
    response = requests.get(url, headers=headers)
    sha = None
    if response.status_code == 200:
        data = response.json()
        sha = data["sha"]
        content = base64.b64decode(data["content"]).decode("utf-8")
    else: content = "\ufeff"
    output = io.StringIO()
    writer = csv.writer(output)
    if content == "\ufeff": writer.writerow(header_data)
    writer.writerow(row_data)
    new_content = content + output.getvalue()
    payload = {"message": f"Added: {filename}", "content": base64.b64encode(new_content.encode("utf-8")).decode("utf-8")}
    if sha: payload["sha"] = sha
    requests.put(url, headers=headers, json=payload)

def log_generation(title_text, authors_text, lang):
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    row = [timestamp, lang, title_text, authors_text]
    header = ["Уақыты (Timestamp)", "Тіл (Language)", "Тақырып (Title)", "Авторлар (Authors)"]
    append_to_github_csv("generation_logs.csv", row, header)

def log_registration(name, email, phone, org, pos):
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    row = [timestamp, name, email, phone, org, pos]
    header = ["Уақыты (Timestamp)", "Аты-жөні (Full Name)", "Email", "Телефон (Phone)", "Ұйым (Organization)", "Лауазымы (Position)"]
    append_to_github_csv("registered_users.csv", row, header)

def convert_to_pdf(docx_path, pdf_path):
    """Оболочка для конвертации DOCX -> PDF через LibreOffice или docx2pdf"""
    try:
        # Для Linux/Серверов с LibreOffice
        subprocess.run(['soffice', '--headless', '--convert-to', 'pdf', docx_path, '--outdir', os.path.dirname(pdf_path)], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        if os.path.exists(pdf_path): return True
    except: pass
    
    try:
        # Запасной вариант вызова (некоторые дистрибутивы Linux)
        subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', docx_path, '--outdir', os.path.dirname(pdf_path)], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        if os.path.exists(pdf_path): return True
    except: pass

    try:
        # Для локальных машин Windows/Mac с установленным MS Word
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        if os.path.exists(pdf_path): return True
    except: pass

    return False

# ------------ Header ------------
hc1, hc2, hc3 = st.columns([6, 1.8, 1.8])
with hc1:
    st.title(l["title"])
    st.caption(l["subtitle"])
with hc2:
    _lang_labels = {"kz": "🇰🇿 Қазақша", "ru": "🇷🇺 Русский", "en": "🇬🇧 English"}
    _lang_keys = list(_lang_labels.keys())
    _sel = st.selectbox("lang", _lang_keys, index=_lang_keys.index(st.session_state.lang), format_func=lambda x: _lang_labels[x], label_visibility="collapsed")
    if _sel != st.session_state.lang:
        st.session_state.lang = _sel
        st.rerun()
with hc3:
    _tbtn = l["btn_theme_light"] if st.session_state.theme == "dark" else l["btn_theme_dark"]
    if st.button(_tbtn, use_container_width=True):
        st.session_state.theme = "light" if st.session_state.theme == "dark" else "dark"
        st.rerun()

st.markdown("---")

if "nav_radio" not in st.session_state or st.session_state.nav_radio not in [l["nav_gen"], l["nav_reg"]]:
    st.session_state.nav_radio = l["nav_gen"]
if st.session_state.get("go_to_gen"):
    st.session_state.nav_radio = l["nav_gen"]
    st.session_state.go_to_gen = False

app_mode = st.radio("", [l["nav_gen"], l["nav_reg"]], horizontal=True, label_visibility="collapsed", key="nav_radio")
st.markdown("---")

is_locked = not st.session_state.is_registered

# ==========================================
# РЕЖИМ: ГЕНЕРАТОР
# ==========================================
if app_mode == l["nav_gen"]:
    if is_locked:
        st.error(l["reg_req_msg"], icon="🔒")

    st.subheader(l["sidebar_title"])
    col_s1, col_s2, col_s3, col_s4, col_s5 = st.columns(5)
    with col_s1: primary_lang = st.selectbox(l["lbl_lang"], ["Русский", "Қазақша", "English"], disabled=is_locked)
    with col_s2: section = st.selectbox(l["lbl_sec"], ["Химия", "География"], disabled=is_locked)
    with col_s3: paper_type = st.selectbox(l["lbl_type"], ["Научная статья (Article)", "Обзор (Review)", "Мини-обзор (Mini-review)", "Краткое сообщение (Communication)"], disabled=is_locked)
    with col_s4: mrnti = st.text_input(l["lbl_mrnti"], value="06.81.23", disabled=is_locked)
    with col_s5: 
        new_font = st.selectbox(l["lbl_ui_font"], list(font_mapping.keys()), index=list(font_mapping.keys()).index(st.session_state.ui_font))
        if new_font != st.session_state.ui_font:
            st.session_state.ui_font = new_font
            st.rerun()

    st.markdown("<br>", unsafe_allow_html=True)

    st.header(l["sec_meta"])
    col1, col2 = st.columns(2)
    with col1:
        title = st.text_area(l["lbl_title"], height=68, disabled=is_locked)
        authors = st.text_area(l["lbl_authors"], help=l["lbl_authors_help"], height=68, disabled=is_locked)
    with col2:
        affiliations = st.text_area(l["lbl_affil"], help=l["lbl_affil_help"], height=68, disabled=is_locked)
        corr_email = st.text_input(l["lbl_email"], disabled=is_locked)

    st.header(l["sec_text"])
    abstract = st.text_area(l["lbl_abstract"], height=150, disabled=is_locked)
    abstract_word_count = len(abstract.split()) if abstract else 0
    if not is_locked:
        if abstract_word_count > 300: st.error(l["err_abs_len"].format(count=abstract_word_count))
        elif abstract_word_count > 0: st.success(l["succ_abs_len"].format(count=abstract_word_count))

    keywords = st.text_input(l["lbl_kw"], help=l["lbl_kw_help"], disabled=is_locked)
    
    st.markdown("##### " + l["lbl_samples"])
    col_dl1, col_dl2, col_dl3, col_dl4, col_dl5 = st.columns(5)
    with col_dl1: st.download_button("📥 Intro", create_sample_docx("Introduction"), file_name="sample_intro.docx", use_container_width=True, disabled=is_locked)
    with col_dl2: st.download_button("📥 Methods", create_sample_docx("Materials and Methods"), file_name="sample_methods.docx", use_container_width=True, disabled=is_locked)
    with col_dl3: st.download_button("📥 Results", create_sample_docx("Results"), file_name="sample_results.docx", use_container_width=True, disabled=is_locked)
    with col_dl4: st.download_button("📥 Discussion", create_sample_docx("Discussion"), file_name="sample_discussion.docx", use_container_width=True, disabled=is_locked)
    with col_dl5: st.download_button("📥 Conclusion", create_sample_docx("Conclusion"), file_name="sample_conclusion.docx", use_container_width=True, disabled=is_locked)
    st.markdown("<br>", unsafe_allow_html=True)

    col_i1, col_i2, col_i3 = st.columns([1,1,1])
    with col_i1:
        file_intro = st.file_uploader(l["lbl_intro"], type=["txt", "docx"], disabled=is_locked)
        file_methods = st.file_uploader(l["lbl_methods"], type=["txt", "docx"], disabled=is_locked)
    with col_i2:
        file_results = st.file_uploader(l["lbl_results"], type=["txt", "docx"], disabled=is_locked)
        file_discussion = st.file_uploader(l["lbl_discussion"], type=["txt", "docx"], disabled=is_locked)
    with col_i3:
        file_conclusion = st.file_uploader(l["lbl_conclusion"], type=["txt", "docx"], disabled=is_locked)

    st.markdown("<br><hr>", unsafe_allow_html=True)
    
    # --- ДИНАМИЧЕСКИЕ МЕНЕДЖЕРЫ РИСУНКОВ И ТАБЛИЦ С КОМПАКТНЫМИ ЗАГРУЗЧИКАМИ ---
    col_ft1, col_ft2 = st.columns(2)
    
    # --- ФОРМА РИСУНКОВ ---
    with col_ft1:
        st.header(l["lbl_fig_manager"])
        with st.expander(l["lbl_fig_hint_title"]):
            st.markdown(l["lbl_fig_hint_text"])
        
        hf1, hf2, hf3 = st.columns([1.5, 3.5, 3])
        hf1.markdown("**Tag**")
        hf2.markdown("**Caption**")
        hf3.markdown("**File Upload**")

        for i in range(st.session_state.fig_count):
            cf1, cf2, cf3 = st.columns([1.5, 3.5, 3])
            with cf1: st.text_input(f"fig_tag_{i}", value=f"[@fig{i+1}]", key=f"f_tag_{i}", label_visibility="collapsed", disabled=is_locked)
            with cf2: st.text_input(f"fig_cap_{i}", placeholder="Caption...", key=f"f_cap_{i}", label_visibility="collapsed", disabled=is_locked)
            with cf3: 
                # Скрытый маркер для применения CSS к компактному загрузчику
                st.markdown('<div class="small-uploader-marker"></div>', unsafe_allow_html=True)
                st.file_uploader(f"fig_file_{i}", type=["png", "jpg", "jpeg"], key=f"f_file_{i}", label_visibility="collapsed", disabled=is_locked)
            
        if st.button(l["lbl_add_fig"], disabled=is_locked):
            st.session_state.fig_count += 1
            st.rerun()

    # --- ФОРМА ТАБЛИЦ ---
    with col_ft2:
        st.header(l["lbl_tab_manager"])
        with st.expander(l["lbl_tab_hint_title"]):
            st.markdown(l["lbl_tab_hint_text"])
        
        ht1, ht2, ht3 = st.columns([1.5, 3.5, 3])
        ht1.markdown("**Tag**")
        ht2.markdown("**Caption**")
        ht3.markdown("**File Upload**")

        for i in range(st.session_state.tab_count):
            ct1, ct2, ct3 = st.columns([1.5, 3.5, 3])
            with ct1: st.text_input(f"tab_tag_{i}", value=f"[@tab{i+1}]", key=f"t_tag_{i}", label_visibility="collapsed", disabled=is_locked)
            with ct2: st.text_input(f"tab_cap_{i}", placeholder="Caption...", key=f"t_cap_{i}", label_visibility="collapsed", disabled=is_locked)
            with ct3: 
                # Скрытый маркер для применения CSS к компактному загрузчику
                st.markdown('<div class="small-uploader-marker"></div>', unsafe_allow_html=True)
                st.file_uploader(f"tab_file_{i}", type=["xlsx", "csv", "docx", "txt"], key=f"t_file_{i}", label_visibility="collapsed", disabled=is_locked)
            
        if st.button(l["lbl_add_tab"], disabled=is_locked):
            st.session_state.tab_count += 1
            st.rerun()

    # --- Әдебиеттер менеджері (Smart Reference Manager) ---
    st.markdown("<hr>", unsafe_allow_html=True)
    st.header(l["lbl_ref_manager"])
    ref_style = st.selectbox(l["lbl_ref_style"], ["GOST", "APA", "IEEE"], disabled=is_locked)
    ref_df = pd.DataFrame([{"Tag in text": "[@ref1]", "Author(s)": "", "Year": "", "Title": "", "Journal/Publisher": "", "Volume/Pages": ""}])
    if not is_locked:
        edited_refs = st.data_editor(ref_df, num_rows="dynamic", use_container_width=True)
    else:
        st.dataframe(ref_df, use_container_width=True)

    # --- Дополнительная информация (Back Matter) ---
    st.header(l["sec_backmatter"])
    val_supp = st.text_area(l["lbl_supp"], value="No supplementary material.", height=68, disabled=is_locked)
    val_contrib = st.text_area(l["lbl_contrib"], value="Conceptualization, X.X. and Y.Y.; methodology, X.X.; software, X.X.; validation, X.X., Y.Y. and Z.Z.; formal analysis, X.X.; investigation, X.X.; resources, X.X.; data curation, X.X.; writing—original draft preparation, X.X.; writing—review and editing, X.X.; visualisation, X.X.; supervision, X.X.; project administration, X.X.; funding acquisition, Y.Y. All authors have read and agreed to the published version of the manuscript.", height=120, disabled=is_locked)
    val_auth_info = st.text_area(l["lbl_auth_info"], value="Beisembayev, Adil Sayatuly - researcher, L.N. Gumilyov Eurasian National University, Kazhymukan st., 13, Astana, Kazakhstan, 010000; email: beisembayev_as@enu.kz, https://orcid.org/0001-0003-2203-9099", height=80, disabled=is_locked)
    val_funding = st.text_area(l["lbl_funding"], value="This research received no external funding.", height=68, disabled=is_locked)
    val_ack = st.text_area(l["lbl_ack"], value="Administrative and technical support was provided by...", height=68, disabled=is_locked)
    val_coi = st.text_area(l["lbl_coi"], value="The authors declare no conflicts of interest. The funders had no role in the study’s design, data collection, analysis, manuscript writing, or publication decisions.", height=80, disabled=is_locked)

    # --- Аудармалар ---
    st.header(l["sec_trans"])
    st.info(l["trans_info"])
    trans_langs = ["Русский", "Қазақша", "English"]
    if primary_lang in trans_langs: trans_langs.remove(primary_lang)

    col_t1, col_t2 = st.columns(2)
    with col_t1:
        st.subheader(f"{trans_langs[0]}")
        t1_title = st.text_input(f"{l['lbl_title']} ({trans_langs[0]})", disabled=is_locked)
        t1_authors = st.text_input(f"{l['lbl_authors']} ({trans_langs[0]})", disabled=is_locked)
        t1_abstract = st.text_area(f"{l['lbl_abstract']} ({trans_langs[0]})", height=100, disabled=is_locked)
        t1_keywords = st.text_input(f"{l['lbl_kw']} ({trans_langs[0]})", disabled=is_locked)
    with col_t2:
        st.subheader(f"{trans_langs[1]}")
        t2_title = st.text_input(f"{l['lbl_title']} ({trans_langs[1]})", disabled=is_locked)
        t2_authors = st.text_input(f"{l['lbl_authors']} ({trans_langs[1]})", disabled=is_locked)
        t2_abstract = st.text_area(f"{l['lbl_abstract']} ({trans_langs[1]})", height=100, disabled=is_locked)
        t2_keywords = st.text_input(f"{l['lbl_kw']} ({trans_langs[1]})", disabled=is_locked)

    st.markdown("---")
    generate_btn = st.button(l["gen_btn"], type="primary", use_container_width=True, disabled=is_locked)

    if generate_btn and not is_locked:
        if abstract_word_count > 300: st.error(l["err_abs_len"].format(count=abstract_word_count))
        elif not title or not authors: st.warning(l["err_fill_req"])
        else:
            with st.spinner("Генерация документов..."):
                try:
                    # 1. Мәтіндерді жинақтау
                    main_text_compiled = ""
                    if file_intro: main_text_compiled += "1. INTRODUCTION\n" + extract_text(file_intro) + "\n\n"
                    if file_methods: main_text_compiled += "2. MATERIALS AND METHODS\n" + extract_text(file_methods) + "\n\n"
                    if file_results: main_text_compiled += "3. RESULTS\n" + extract_text(file_results) + "\n\n"
                    if file_discussion: main_text_compiled += "4. DISCUSSION\n" + extract_text(file_discussion) + "\n\n"
                    if file_conclusion: main_text_compiled += "5. CONCLUSION\n" + extract_text(file_conclusion) + "\n\n"
                    
                    # 2. Обработка динамических Рисунков
                    fig_text_compiled = ""
                    fig_counter = 1
                    for i in range(st.session_state.fig_count):
                        c_tag = st.session_state.get(f"f_tag_{i}", "").strip()
                        c_cap = st.session_state.get(f"f_cap_{i}", "").strip()
                        
                        if c_cap:
                            fig_label = f"{l['fig_prefix']} {fig_counter}"
                            fig_text_compiled += f"{fig_label}. {c_cap}\n"
                            if c_tag: main_text_compiled = main_text_compiled.replace(c_tag, fig_label)
                            fig_counter += 1

                    # 3. Обработка динамических Таблиц
                    tab_text_compiled = ""
                    tab_counter = 1
                    for i in range(st.session_state.tab_count):
                        c_tag = st.session_state.get(f"t_tag_{i}", "").strip()
                        c_cap = st.session_state.get(f"t_cap_{i}", "").strip()
                        
                        if c_cap:
                            tab_label = f"{l['tab_prefix']} {tab_counter}"
                            tab_text_compiled += f"{tab_label}. {c_cap}\n"
                            if c_tag: main_text_compiled = main_text_compiled.replace(c_tag, tab_label)
                            tab_counter += 1
                    
                    if fig_text_compiled or tab_text_compiled:
                        main_text_compiled += "\n\n--- FIGURES & TABLES ---\n" + fig_text_compiled + "\n" + tab_text_compiled

                    # 4. Back Matter
                    back_matter = ""
                    if val_supp: back_matter += f"6. Supplementary Materials\n{val_supp}\n\n"
                    if val_contrib: back_matter += f"7. Author Contributions\n{val_contrib}\n\n"
                    if val_auth_info: back_matter += f"8. Author Information\n{val_auth_info}\n\n"
                    if val_funding: back_matter += f"9. Funding\n{val_funding}\n\n"
                    if val_ack: back_matter += f"10. Acknowledgements\n{val_ack}\n\n"
                    if val_coi: back_matter += f"11. Conflicts of Interest\n{val_coi}\n\n"
                    main_text_compiled += "\n\n" + back_matter

                    # 5. References
                    refs_compiled = []
                    ref_counter = 1
                    for _, row in edited_refs.iterrows():
                        r_tag = str(row.get("Tag in text", "")).strip()
                        r_author = str(row.get("Author(s)", "")).strip()
                        r_year = str(row.get("Year", "")).strip()
                        r_title = str(row.get("Title", "")).strip()
                        r_journal = str(row.get("Journal/Publisher", "")).strip()
                        r_vol = str(row.get("Volume/Pages", "")).strip()
                        
                        if r_author == "nan" or not r_author: continue
                        
                        if ref_style == "APA":
                            ref_entry = f"{r_author} ({r_year}). {r_title}. {r_journal}, {r_vol}."
                            first_author = r_author.split(',')[0].strip()
                            in_text_citation = f"({first_author} et al., {r_year})"
                        elif ref_style == "IEEE":
                            ref_entry = f"[{ref_counter}] {r_author}, \"{r_title},\" {r_journal}, {r_vol}, {r_year}."
                            in_text_citation = f"[{ref_counter}]"
                        else: # GOST
                            ref_entry = f"{ref_counter}. {r_author} {r_title} // {r_journal}. - {r_year}. - {r_vol}."
                            in_text_citation = f"[{ref_counter}]"
                            
                        refs_compiled.append(ref_entry)
                        if r_tag and r_tag != "nan": main_text_compiled = main_text_compiled.replace(r_tag, in_text_citation)
                        ref_counter += 1
                    final_references = "\n".join(refs_compiled)

                    # Шаблон
                    template_filename = "Russian_template_2025.docx"
                    if primary_lang == "Русский": template_filename = "Russian_template_2025.docx"
                    elif primary_lang == "Қазақша": template_filename = "Kazakh_template_2025.docx"
                    elif primary_lang == "English": template_filename = "English_template_2025.docx"

                    context = {
                        "mrnti": mrnti, "section": section, "paper_type": paper_type,
                        "title": title, "authors": authors, "affiliations": affiliations, "corr_email": corr_email,
                        "abstract": abstract, "keywords": keywords,
                        "main_text": main_text_compiled, "references": final_references,
                        "t1_title": t1_title, "t1_authors": t1_authors, "t1_abstract": t1_abstract, "t1_keywords": t1_keywords,
                        "t2_title": t2_title, "t2_authors": t2_authors, "t2_abstract": t2_abstract, "t2_keywords": t2_keywords,
                    }

                    doc = DocxTemplate(template_filename)
                    doc.render(context)
                    
                    # Генерация файлов через TemporaryDirectory
                    with tempfile.TemporaryDirectory() as tmpdir:
                        docx_path = os.path.join(tmpdir, "Formatted_Article.docx")
                        pdf_path = os.path.join(tmpdir, "Formatted_Article.pdf")
                        
                        doc.save(docx_path)
                        with open(docx_path, "rb") as f:
                            docx_bytes = f.read()

                        pdf_success = convert_to_pdf(docx_path, pdf_path)
                        pdf_bytes = None
                        if pdf_success:
                            with open(pdf_path, "rb") as f:
                                pdf_bytes = f.read()
                    
                    st.success(l["succ_gen"])
                    log_generation(title, authors, primary_lang)

                    dcol1, dcol2 = st.columns(2)
                    with dcol1:
                        st.download_button(label=l["btn_dl_docx"], data=docx_bytes, file_name="Formatted_Article.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="primary", use_container_width=True)
                    with dcol2:
                        if pdf_bytes:
                            st.download_button(label=l["btn_dl_pdf"], data=pdf_bytes, file_name="Formatted_Article.pdf", mime="application/pdf", type="primary", use_container_width=True)
                        else:
                            st.warning(l["err_pdf"])
                            
                except Exception as e:
                    st.error(f"{l['err_gen']} {e}")
                    st.info("💡 Ескерту: 'Russian_template_2025.docx', 'Kazakh_template_2025.docx' және 'English_template_2025.docx' файлдары бумада болуы тиіс.")

# ==========================================
# РЕЖИМ: РЕГИСТРАЦИЯ
# ==========================================
elif app_mode == l["nav_reg"]:
    st.header(l["reg_header"])
    if st.session_state.is_registered:
        st.success(l["reg_success"])
        st.info(l["reg_info"])
    else:
        with st.form("registration_form"):
            r_name = st.text_input(l["reg_name"])
            r_email = st.text_input(l["reg_email"])
            r_phone = st.text_input(l["reg_phone"])
            r_org = st.text_input(l["reg_org"])
            r_pos = st.text_input(l["reg_pos"])
            submitted = st.form_submit_button(l["reg_submit"], type="primary")

            if submitted:
                if r_name and r_email and r_phone:
                    with st.spinner("Тіркелу жүріп жатыр..."):
                        log_registration(r_name, r_email, r_phone, r_org, r_pos)
                    st.session_state.is_registered = True
                    st.session_state.go_to_gen = True
                    st.success(l["reg_success"])
                    st.rerun()
                else: st.error(l["reg_err_fill"])

with st.sidebar:
    if os.path.exists("generation_logs.csv") or os.path.exists("registered_users.csv"):
        st.markdown("<br><br><br>", unsafe_allow_html=True)
        st.caption("🔒 Панель администратора")
        if os.path.exists("generation_logs.csv"):
            with open("generation_logs.csv", "rb") as f:
                st.download_button(label="📊 Логи генерации (.csv)", data=f, file_name="generation_logs.csv", mime="text/csv", use_container_width=True)
        if os.path.exists("registered_users.csv"):
            with open("registered_users.csv", "rb") as f:
                st.download_button(label="👥 База пользователей (.csv)", data=f, file_name="registered_users.csv", mime="text/csv", use_container_width=True)

st.markdown("---")
st.markdown(
    f'<div style="text-align:center;font-size:12px;color:gray;padding:12px 0 20px 0;line-height:2.2;">'
    f'<b style="font-size:13px;">© 2025 {l["f_author"]}</b><br>'
    f'📧 <a href="mailto:samarkhanov_kb@enu.kz" style="text-decoration:none;">samarkhanov_kb@enu.kz</a>'
    f'&nbsp;·&nbsp;<a href="mailto:kanat.baurzhanuly@gmail.com" style="text-decoration:none;">kanat.baurzhanuly@gmail.com</a><br>'
    f'🏛️ <a href="https://fns.enu.kz/kz/page/departments/physical-and-economical-geography/faculty-members" target="_blank" style="text-decoration:none;">{l["f_univ"]}</a><br>'
    f'</div>', unsafe_allow_html=True)
